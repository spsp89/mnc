from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "proposal_assets"
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "detailed-seo-aeo-uiux-report-chams-ecosystem.docx"

SCORECARD = OUT_DIR / "audit-scorecard.png"
ROADMAP = OUT_DIR / "roadmap-90-days.png"
ROI = OUT_DIR / "roi-scenarios.png"

NAVY = RGBColor(11, 47, 116)
NAVY_DEEP = RGBColor(4, 28, 85)
BLUE = RGBColor(30, 92, 179)
GOLD = RGBColor(244, 178, 39)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(97, 113, 133)


SITES = [
    {
        "name": "CHAMS Global",
        "url": "https://chamsglobal.com/",
        "screenshot": "chamsglobal.png",
        "role": "Master corporate brand, service authority hub, and trust-transfer layer for the ecosystem.",
        "scores": {"UI": 6, "UX": 5, "SEO": 5, "AEO": 3, "SMO": 6, "CRO": 5},
        "headline": "Well-positioned brand platform with decent breadth, but the site needs stronger proof architecture, cleaner metadata consistency, and better answer-engine formatting.",
        "strengths": [
            "Broad service coverage and clear agency positioning across Kerala and GCC-oriented markets.",
            "Blog and service landing pages create a usable SEO foundation.",
            "The brand already has enough authority signals to become the ecosystem’s main conversion and case-study hub.",
        ],
        "uiux": [
            "The homepage looks credible but the hero is visually light on proof and heavy on generic positioning language.",
            "The journey from awareness to trust to conversion is not sequenced tightly enough. Visitors see services, but outcomes, case studies, and reasons to believe are underplayed.",
            "The site would benefit from clearer service page hierarchy, stronger internal CTA logic, and more obvious lead capture modules."
        ],
        "technical": [
            ["Homepage canonical points to `https://www.chamsglobal.com/` while the reviewed domain resolved at `https://chamsglobal.com/`.", "Canonical normalization should be cleaned up to reduce ambiguity around the preferred host."],
            ["Homepage had no visible Open Graph title, description, or image in the metadata review.", "Social sharing previews are weaker than they should be, reducing click-through from WhatsApp, LinkedIn, and other social surfaces."],
            ["Homepage schema detection surfaced `BreadcrumbList` only.", "The site is underusing structured data such as Organization, LocalBusiness, Service, FAQ, Person, and WebSite schema."],
            ["Sample page `digital-marketing-agency-in-dubai` showed duplicate H1 presence.", "Heading hierarchy should be tightened so pages present one primary entity/topic cleanly."],
            ["The sitemap returned 60 URLs and several keyword-targeted service pages.", "Good foundation, but content quality, uniqueness, and interlinking need strategic consolidation."],
        ],
        "page_samples": [
            ["Homepage", "Branding Agency in Kochi | Marketing Agency in Kerala", "4 H1s detected", "Too many H1s on the homepage can dilute topical clarity."],
            ["About", "Top Advertising Agency in Kerala | Chamsglobal", "2 H1s", "Multiple H1s again suggest structural cleanup is needed."],
            ["Blog", "Updated Marketing Blogs by the Top Media Agency in Kerala", "4455 words", "Good depth, but editorial structure and topic clustering can be stronger."],
            ["Branding page", "Branding Agency in Kochi | Best Advertising Agency in Kochi", "736 words", "Service landing page has useful depth and can be further upgraded for local intent."],
            ["SMM page", "Social Media Marketing | Digital Marketing Agency in Kochi", "1289 words", "Strong opportunity to reformat into problem-solution-result-FAQ sections for AEO."],
        ],
        "seo_findings": [
            "The domain already ranks/indexes for multiple service terms, which is valuable.",
            "Metadata is present on many pages, but social metadata and schema are incomplete.",
            "The site needs page-type discipline: homepage, service page, location page, case study, blog article, and founder page should follow clearer templates.",
            "Local intent pages like branding and marketing in Kochi/Kerala are promising, but should link more strongly to proof assets and enquiry intent."
        ],
        "aeo_findings": [
            "The site is not yet shaped for answer engines. Many pages read like marketing copy rather than answer-ready expert content.",
            "There are few obvious FAQ structures, comparison blocks, pricing guidance, or process explainers on key pages.",
            "AEO opportunity is especially strong around strategic topics such as branding, agency selection, marketing planning, OOH, and digital strategy."
        ],
        "smo_findings": [
            "Social links exist, but some are low-quality destinations like a generic Twitter home link or admin-style LinkedIn endpoints.",
            "The site should be supported by a stronger loop: case study snippet -> social proof -> landing page -> enquiry form.",
            "CHAMS is the natural home for polished case-study cards and expert founder clips."
        ],
        "business_potential": [
            "Highest value as the ecosystem’s authority and conversion hub.",
            "Can attract premium B2B clients if proof, trust, and strategic messaging become more concrete.",
            "Should become the main case-study destination feeding leads into CHAMS, KSRTC, and Let’s Play Outdoors offers."
        ],
        "priority": [
            "Rebuild homepage above-the-fold with clearer offer segmentation, proof badges, and CTA logic.",
            "Add Organization, LocalBusiness, Service, FAQ, and WebSite schema across priority templates.",
            "Create better case-study, testimonial, and founder-thought-leadership modules.",
            "Normalize canonical and host handling across www vs non-www.",
        ],
        "keywords": [
            "branding agency in Kochi",
            "marketing agency in Kerala",
            "digital marketing agency in Dubai",
            "advertising agency in Kerala",
            "brand promotion agency in Kochi",
        ],
    },
    {
        "name": "KSRTC Bus Branding",
        "url": "https://www.ksrtcbusbranding.com/",
        "screenshot": "ksrtc_retry.jpg",
        "role": "Specialized niche acquisition property for transit-media and bus-branding intent.",
        "scores": {"UI": 6, "UX": 5, "SEO": 4, "AEO": 2, "SMO": 5, "CRO": 6},
        "headline": "Commercially attractive niche site with clear offer visibility, but its technical hygiene and information architecture need significant strengthening.",
        "strengths": [
            "The homepage communicates the offer clearly and quickly.",
            "The site is focused on a specific niche with real purchase intent.",
            "Gallery-led assets are useful trust builders when properly optimized."
        ],
        "uiux": [
            "The visual proposition is strong, but the site feels older and less polished than the opportunity deserves.",
            "The first impression is product-clear, but conversion UX is not modernized enough for a high-value B2B enquiry journey.",
            "There is room for stronger quote-request framing, inventory explanation, route coverage detail, and campaign process guidance."
        ],
        "technical": [
            ["Homepage returned no H1 during review.", "A missing H1 on the main page weakens topical clarity for search engines and screen readers."],
            ["Both `/` and `/index.php` were live and indexable with the same title and content.", "This creates duplication risk and should be canonicalized or redirected."],
            ["No homepage JSON-LD schema was detected.", "The site is missing key opportunities for Organization, Service, FAQ, VideoObject, and Breadcrumb schema."],
            ["`/admin/`, `/admin/uploads/`, and `/admin/uploads/gallery/` were publicly reachable.", "This is a trust, SEO, and operational hygiene issue. Sensitive or irrelevant directories should not be publicly indexable."],
            ["Robots file appears template-generated.", "Robots and indexing policy should be reviewed as part of a proper technical SEO cleanup."],
        ],
        "page_samples": [
            ["Homepage", "KSRTC Bus Branding | KSRTC Bus Advertising Agency in Kerala", "0 H1", "Important structural SEO issue on the main page."],
            ["About", "KSRTC Advertising Agency Kerala | Bus Branding", "594 words", "Solid trust page, but could be modernized and strengthened with proof."],
            ["Advantages", "Low Cost Bus Advertising | Ad On Bus Cost", "307 words", "Good commercial-intent page that can be expanded with FAQs and pricing logic."],
            ["Gallery pages", "Volvo / Low Floor / KSRTC Gallery", "134 to 137 words", "Useful for trust, but image alt text, captions, and internal links should be stronger."],
            ["Video gallery", "Transit Advertising Videos", "indexed snippet visible", "Can support richer SERP visibility if VideoObject schema is added."],
        ],
        "seo_findings": [
            "This site has niche-keyword relevance and can become a strong local-search performer for transit-media queries.",
            "It needs better content depth on buying questions: cost, route coverage, ad sizes, booking process, timelines, creative support, and reporting.",
            "Page duplication and public-directory exposure reduce technical quality.",
            "Gallery and video assets can support richer search results if properly structured."
        ],
        "aeo_findings": [
            "The site currently says what it offers, but not in a way that answer engines can reuse easily.",
            "It should contain FAQ and explainer blocks such as: What is KSRTC bus branding, how much does it cost, where are the routes, what is the campaign duration, what creatives work best, and how to measure ROI.",
            "This site is an ideal candidate for concise commercial-answer content because the niche is narrow and buyer questions are predictable."
        ],
        "smo_findings": [
            "The business has visual media assets that should perform well on Instagram, YouTube Shorts, LinkedIn, and WhatsApp circulation.",
            "Needs more campaign stories, route visibility clips, mockup-to-live transformations, and client success explanations."
        ],
        "business_potential": [
            "Very strong short-term lead potential because the offer is specific and commercially meaningful.",
            "With stronger conversion paths and page trust, this site can become a consistent niche lead source."
        ],
        "priority": [
            "Fix homepage H1 and duplication between `/` and `/index.php`.",
            "Review and secure public admin/upload paths.",
            "Add Service, FAQ, Organization, Breadcrumb, and VideoObject schema.",
            "Create commercial landing pages for pricing, route coverage, and campaign process."
        ],
        "keywords": [
            "KSRTC bus branding",
            "bus advertising agency in Kerala",
            "transit advertising Kerala",
            "bus branding cost Kerala",
            "Volvo bus advertising Kerala",
        ],
    },
    {
        "name": "Let's Play Outdoors",
        "url": "https://letsplayoutdoors.in/",
        "screenshot": "lpo_retry.jpg",
        "role": "Marketplace-scale lead engine and inventory discovery platform for OOH/DOOH products.",
        "scores": {"UI": 7, "UX": 6, "SEO": 6, "AEO": 4, "SMO": 6, "CRO": 7},
        "headline": "The strongest long-term growth asset in the group, with real inventory-led SEO potential, but it needs cleaner index discipline and better landing-page architecture.",
        "strengths": [
            "The search-first marketplace concept is commercially strong and scalable.",
            "Homepage intent is clear and the platform is differentiated from a normal agency site.",
            "Open Graph setup is stronger here than on the other reviewed sites."
        ],
        "uiux": [
            "The homepage communicates an online-booking proposition effectively.",
            "The search interface is useful, but the page still feels broad and somewhat directory-like.",
            "It needs stronger category access, trust modules, featured collections, and conversion framing to reduce decision friction."
        ],
        "technical": [
            ["Homepage canonical points to `/home` rather than the root URL.", "Canonical setup should be reviewed so the preferred public URL structure is consistent."],
            ["Contact page returned 0 detected H1s and only ~79 words.", "Important utility pages need stronger content and proper structural headings."],
            ["Blog listing page returned 0 H1s.", "Blog index pages should still carry proper structural signals."],
            ["FAQ page returned 0 H1s.", "A core answer-engine page should have explicit heading structure."],
            ["At least two sample URLs from search/sitemap review returned empty metadata or zero-word-like output.", "Potential thin/empty pages should be repaired, redirected, or removed from indexation."],
        ],
        "page_samples": [
            ["Homepage /home", "Outdoor Advertising Agencies | Hoarding Companies | Kochi, Kerala", "785 words", "Healthy core page, but category entry and trust hierarchy can improve."],
            ["About us", "About Us - Outdoor Advertising Agencies", "search result visible", "Good positioning page and worth expanding for entity clarity."],
            ["Blog listing", "Blog | Lets play outdoors", "208 words, 0 H1", "Index page needs structural cleanup."],
            ["FAQ", "FAQ", "519 words, 0 H1", "Good foundation for AEO if restructured and schematized."],
            ["OOH article", "Walking Billboards – How It Can Change Your Advertising Game", "743 words", "Useful content pattern for long-tail SEO."],
        ],
        "seo_findings": [
            "This domain has the biggest programmatic and category-SEO upside of the group.",
            "Its sitemap footprint is large enough to support a serious content and category strategy.",
            "Index quality must be tightened so empty or broken pages do not dilute the site.",
            "City pages, format pages, product pages, and pricing-guidance pages can create durable search traffic."
        ],
        "aeo_findings": [
            "The domain is well-suited to answer-intent content because buyers often ask repeated questions about OOH formats, city availability, costs, lead times, and suitability by objective.",
            "FAQ exists, which is a positive sign, but it should be restructured into explicit question-answer modules and enhanced with FAQ schema.",
            "Comparison pages such as billboard vs bus branding vs DOOH vs roadshow vehicle can perform well for AI summaries and commercial search."
        ],
        "smo_findings": [
            "This brand is inherently visual and can generate strong social assets from inventory highlights and city/location-based showcases.",
            "Social content should send traffic to themed landing pages rather than only to the homepage."
        ],
        "business_potential": [
            "Highest long-term traffic and scale potential in the ecosystem.",
            "Can become the top-of-funnel engine for both direct bookings and assisted-sales leads."
        ],
        "priority": [
            "Fix index quality on empty or thin pages.",
            "Create stronger category, city, and format landing pages.",
            "Restructure FAQ and buying guides for AEO.",
            "Improve contact, blog, and FAQ template H1 hierarchy."
        ],
        "keywords": [
            "outdoor advertising in Kerala",
            "OOH advertising marketplace India",
            "billboard advertising Kerala",
            "bus branding Kerala",
            "DOOH advertising India",
        ],
    },
    {
        "name": "Misbah Salam",
        "url": "https://misbahsalam.com/",
        "screenshot": "misbah_retry.jpg",
        "role": "Founder-led authority layer that can improve trust, thought leadership, and entity relevance for the wider group.",
        "scores": {"UI": 7, "UX": 6, "SEO": 5, "AEO": 3, "SMO": 5, "CRO": 5},
        "headline": "Strong personal-brand asset with high authority upside, but it needs content discipline, better metadata differentiation, and clearer consulting-package conversion design.",
        "strengths": [
            "Visually premium and personality-led.",
            "Founder expertise gives the ecosystem a human trust anchor.",
            "Blog-led thought leadership can support both direct consulting and halo-brand visibility."
        ],
        "uiux": [
            "The homepage look is premium, but text formatting and spacing issues reduce perceived polish.",
            "The offer architecture is not yet clear enough for consulting-package conversion.",
            "There should be clearer service pillars, speaking/consulting positioning, and proof-led pathways."
        ],
        "technical": [
            ["Homepage had no visible Open Graph title/description/image in the metadata review.", "Social share previews are underpowered."],
            ["No homepage JSON-LD schema was detected.", "Missing opportunities for Person, Organization, Article, Breadcrumb, FAQ, and WebSite schema."],
            ["Service page returned 0 detected H1s.", "Service pages need stronger heading hierarchy."],
            ["Blog listing page shared the same generic title pattern as the service page.", "Metadata differentiation across page types is weak and should be improved."],
            ["Search snippets visibly revealed awkward spacing and acrostic-style formatting on the homepage biography block.", "Copy presentation and page output need editorial tightening for credibility."],
        ],
        "page_samples": [
            ["Homepage", "Misbah Salam: Strategic Marketing Expert and Branding Specialist...", "1 H1", "Strong main entity, but formatting quality needs improvement."],
            ["Service", "Misbah Salam: Marketing Specialist, Branding, Ads, Strategy...", "0 H1", "Page title too generic and likely reused across multiple pages."],
            ["Blog", "Misbah Salam: Marketing Specialist, Branding, Ads, Strategy...", "0 H1", "Blog index should not share generic metadata with service pages."],
            ["Role of a Brand Marketing Specialist", "Indexed article snippet visible", "Published ~5 months ago", "Good authority-content direction that can be expanded into clusters."],
        ],
        "seo_findings": [
            "The site has real thought-leadership potential, especially around brand strategy, consulting, startup guidance, and leadership positioning.",
            "Metadata differentiation is weak on important templates.",
            "The site should build topic clusters, not just isolated blog posts."
        ],
        "aeo_findings": [
            "This is the best candidate in the ecosystem for expert-answer content, glossary-style explainers, and founder insights.",
            "AEO growth can come from questions like what a brand strategist does, when a business needs rebranding, how to build brand positioning, and how marketing differs across B2B vs B2C."
        ],
        "smo_findings": [
            "Ideal for LinkedIn-first authority, then repurposed to Shorts, Reels, and quote graphics.",
            "Needs stronger clip-to-site routing and recurring editorial themes."
        ],
        "business_potential": [
            "Can materially improve trust transfer to all other websites in the ecosystem.",
            "Can also become a standalone consulting and speaking lead engine if packaging improves."
        ],
        "priority": [
            "Fix metadata differentiation and heading hierarchy across page templates.",
            "Rewrite and polish homepage/service copy for premium credibility.",
            "Build founder-led authority pages, consulting packages, and expert FAQs.",
            "Add Person, Article, FAQ, and Organization schema."
        ],
        "keywords": [
            "brand strategist in Kerala",
            "marketing consultant in Kochi",
            "business strategist Kerala",
            "brand consultant for startups",
            "marketing strategy consultant India",
        ],
    },
]


def set_page_margins(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def set_run_style(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, size=11, color=INK, bold=False, italic=False, after=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    p.alignment = align
    run = p.add_run(text)
    set_run_style(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_title(doc, text, level=1):
    size = {1: 24, 2: 18, 3: 15}.get(level, 13)
    after = {1: 8, 2: 6, 3: 4}.get(level, 4)
    add_para(doc, text, size=size, color=NAVY_DEEP, bold=True, after=after)


def add_bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_style(run, size=10.5)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell, text: str, *, bold=False, color=INK, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    p.alignment = align
    run = p.add_run(text)
    set_run_style(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_simple_table(doc, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, head in zip(table.rows[0].cells, headers):
        shade_cell(cell, "EAF1FB")
        write_cell(cell, head, bold=True, color=NAVY, size=10)
    for row_data in rows:
        row = table.add_row().cells
        for cell, value in zip(row, row_data):
            write_cell(cell, value, size=9.4)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    return table


def build():
    doc = Document()
    for section in doc.sections:
        set_page_margins(section)

    add_para(doc, "Prepared for CHAMS ecosystem websites", size=11, color=MUTED, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Detailed SEO, AEO and UI/UX Audit Report", size=24, color=NAVY_DEEP, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Reviewed domains: chamsglobal.com, ksrtcbusbranding.com, letsplayoutdoors.in, misbahsalam.com", size=11, color=MUTED, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Audit date: 6 July 2026", size=11, color=MUTED, after=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "This version narrows the earlier proposal to the requested scope only: SEO, AEO, and UI/UX improvements. It includes page-level structural findings, content and technical observations, answer-engine opportunities, implementation priorities, and a Kerala-friendly single-package recommendation.",
        size=11,
        after=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_margins(doc.sections[-1])

    add_title(doc, "1. Executive Summary", level=2)
    add_para(
        doc,
        "The CHAMS ecosystem already has the ingredients for meaningful digital growth: a master agency brand, a niche transit-media property, a scalable OOH marketplace, and a founder-led authority site. The issue is not the absence of market fit. The issue is that the digital assets are not yet behaving like a coordinated commercial system.",
        after=8,
    )
    add_para(
        doc,
        "At the moment, each site has some strengths, but the ecosystem loses momentum because of inconsistent metadata quality, incomplete schema, weak AEO structuring, thin page architecture, and uneven conversion design. When fixed, these sites can support each other: CHAMS can close trust gaps, KSRTC can win niche intent, Let’s Play Outdoors can scale organic acquisition, and Misbah Salam can strengthen authority and expert relevance.",
        after=8,
    )
    add_para(doc, "Top conclusions:", bold=True, after=4)
    add_bullets(
        doc,
        [
            "CHAMS Global should become the ecosystem’s main authority, proof, and conversion hub.",
            "KSRTC Bus Branding should be treated as a specialized commercial landing ecosystem, not just a brochure website.",
            "Let’s Play Outdoors has the strongest SEO and marketplace upside, especially for category, city, and format-driven discovery.",
            "Misbah Salam can become the founder-expert layer that lifts trust and answer-engine relevance across the entire group.",
            "The immediate scope should prioritize technical cleanup, stronger commercial page architecture, schema, FAQ structures, and analytics discipline.",
        ],
    )

    add_title(doc, "2. Audit Methodology", level=2)
    add_bullets(
        doc,
        [
            "Live homepage and selected internal-page review performed on 6 July 2026.",
            "Search-result and indexed-footprint checks using branded and `site:`-based query review.",
            "Technical checks included titles, meta descriptions, canonical signals, H1 presence, robots/sitemap availability, structured data presence, and selected directory/index exposure.",
            "Page-level sampling included service pages, blog listings, contact pages, FAQ pages, and gallery pages where available.",
            "Commercial interpretation focused on lead generation, trust transfer, conversion readiness, and long-term organic growth potential.",
        ],
    )

    add_title(doc, "3. Ecosystem Role Mapping", level=2)
    ecosystem_rows = [
        ["CHAMS Global", "Master brand", "Authority, services, trust, case studies", "Premium B2B service enquiries"],
        ["KSRTC Bus Branding", "Niche acquisition site", "Transit-media demand capture", "Focused commercial leads"],
        ["Let's Play Outdoors", "Marketplace growth engine", "Inventory discovery and category SEO", "Scalable lead flow and assisted bookings"],
        ["Misbah Salam", "Founder authority site", "Thought leadership and trust transfer", "Consulting and halo-brand value"],
    ]
    add_simple_table(doc, ["Site", "Role", "Primary purpose", "Business outcome"], ecosystem_rows)
    add_para(
        doc,
        "The strongest future state is not four separate websites competing for attention. It is one ecosystem where each domain has a clear job and passes trust and intent to the others.",
        size=10.5,
        color=MUTED,
        after=8,
    )

    add_title(doc, "4. Scorecard Overview", level=2)
    score_rows = []
    for site in SITES:
        score_rows.append([
            site["name"],
            str(site["scores"]["UI"]),
            str(site["scores"]["UX"]),
            str(site["scores"]["SEO"]),
            str(site["scores"]["AEO"]),
            str(site["scores"]["CRO"]),
        ])
    add_simple_table(doc, ["Site", "UI", "UX", "SEO", "AEO", "CRO"], score_rows)

    for idx, site in enumerate(SITES, start=5):
        doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_margins(doc.sections[-1])
        add_title(doc, f"{idx}. {site['name']}", level=2)
        add_para(doc, site["url"], size=10.5, color=BLUE, after=8)
        add_para(doc, f"Strategic role: {site['role']}", size=11, color=MUTED, bold=True, after=6)
        add_para(doc, site["headline"], size=11, italic=True, after=10)
        shot = ASSET_DIR / site["screenshot"]
        if shot.exists():
            doc.add_picture(str(shot), width=Inches(5.9))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        add_title(doc, "Strengths", level=3)
        add_bullets(doc, site["strengths"])

        add_title(doc, "UI / UX Review", level=3)
        add_bullets(doc, site["uiux"])

        add_title(doc, "Technical and Structural Findings", level=3)
        tech_rows = [[left, right] for left, right in site["technical"]]
        add_simple_table(doc, ["Observation", "Why it matters"], tech_rows, widths=[2.8, 3.5])

        add_title(doc, "Sample Page Review", level=3)
        add_simple_table(
            doc,
            ["Page sample", "Observed title / signal", "Evidence", "Interpretation"],
            site["page_samples"],
            widths=[1.5, 2.0, 1.1, 1.9],
        )

        add_title(doc, "SEO Analysis", level=3)
        add_bullets(doc, site["seo_findings"])

        add_title(doc, "AEO Analysis", level=3)
        add_bullets(doc, site["aeo_findings"])

        add_title(doc, "Business Potential", level=3)
        add_bullets(doc, site["business_potential"])

        add_title(doc, "Priority Actions", level=3)
        add_bullets(doc, site["priority"])

        add_title(doc, "Keyword / Topic Opportunities", level=3)
        add_bullets(doc, site["keywords"])

    doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_margins(doc.sections[-1])
    add_title(doc, "9. Cross-Site Technical Issues", level=2)
    add_bullets(
        doc,
        [
            "Structured data usage is inconsistent and generally underdeveloped across the ecosystem.",
            "Important pages on KSRTC, Let’s Play Outdoors, and Misbah Salam showed missing H1s during sampling.",
            "KSRTC has both `/` and `/index.php` live, indicating canonical/duplication cleanup is needed.",
            "Let’s Play Outdoors showed evidence of thin or empty indexed URLs, which can dilute search quality if not cleaned.",
            "Misbah Salam has weak metadata differentiation across page types.",
            "KSRTC public admin/upload path accessibility should be reviewed immediately from both SEO and operational trust standpoints."
        ],
    )

    add_title(doc, "10. Content Architecture Recommendations", level=2)
    add_para(doc, "Recommended content model by domain:", bold=True, after=4)
    content_rows = [
        ["CHAMS Global", "Service pages, case studies, founder insights, proof pages", "Authority and conversion"],
        ["KSRTC Bus Branding", "Buying guides, route coverage pages, FAQ, pricing explainers, campaign showcases", "High-intent conversion capture"],
        ["Let's Play Outdoors", "Category pages, city pages, format pages, comparisons, FAQs, inventory explainers", "Scalable organic acquisition"],
        ["Misbah Salam", "Expert articles, founder FAQs, consulting package pages, interviews, opinion essays", "Entity authority and trust transfer"],
    ]
    add_simple_table(doc, ["Site", "Content priority", "Primary purpose"], content_rows)
    add_para(doc, "Suggested cluster themes:", bold=True, after=4)
    add_bullets(
        doc,
        [
            "Agency selection: how to choose a branding or marketing partner.",
            "OOH buying education: bus branding, billboards, DOOH, roadshow vehicles, and retail activations.",
            "Commercial intent: cost, duration, city availability, ROI expectations, timelines, and creative requirements.",
            "Founder authority: strategy frameworks, brand positioning, business growth, and market-behavior insights.",
        ],
    )

    add_title(doc, "11. AEO and Schema Roadmap", level=2)
    add_bullets(
        doc,
        [
            "Create FAQ blocks on every priority commercial page and implement FAQ schema where appropriate.",
            "Use Organization / LocalBusiness schema for CHAMS and Let’s Play Outdoors, and relevant Service schema for commercial service pages.",
            "Add Person schema and Article schema to Misbah Salam to strengthen expert identity.",
            "Add VideoObject schema to KSRTC video content and relevant gallery/video assets where technically feasible.",
            "Introduce answer-ready structures: question headings, concise definitions, comparison tables, process sections, and summary bullets.",
        ],
    )

    add_title(doc, "12. Measurement Framework", level=2)
    add_bullets(
        doc,
        [
            "Track leads by source, landing page, domain, and campaign theme.",
            "Separate raw enquiries from qualified enquiries and booked calls.",
            "Measure top commercial pages for scroll depth, CTA clicks, and form submissions.",
            "Track branded vs non-branded search growth and the contribution of each site.",
            "Review lead quality monthly and feed learnings back into content and page design."
        ],
    )

    add_title(doc, "13. Detailed 90-Day Execution Plan", level=2)
    add_para(doc, "Phase 1: Days 1-20", bold=True, after=4)
    add_bullets(
        doc,
        [
            "Fix metadata gaps, H1 issues, canonical inconsistencies, and index-quality problems.",
            "Tighten homepage messaging and CTA logic on CHAMS, KSRTC, and Misbah Salam.",
            "Set analytics, event tracking, and domain-level KPI reporting."
        ],
    )
    add_para(doc, "Phase 2: Days 21-45", bold=True, after=4)
    add_bullets(
        doc,
        [
            "Deploy schema across core templates.",
            "Build FAQ, process, pricing-guidance, and category-support content.",
            "Launch improved service and landing-page architecture for CHAMS and KSRTC.",
            "Clean and enrich Let’s Play Outdoors category and buying-flow pages."
        ],
    )
    add_para(doc, "Phase 3: Days 46-75", bold=True, after=4)
    add_bullets(
        doc,
        [
            "Publish founder-led expert content and cross-link it intelligently into the service ecosystem.",
            "Expand long-tail content clusters and internal linking.",
            "Review enquiry quality and refine pages based on actual commercial behavior."
        ],
    )

    add_title(doc, "14. Deliverables Matrix", level=2)
    deliverables = [
        ["Technical SEO", "Meta cleanup, H1 cleanup, canonical review, sitemap/index discipline, schema deployment"],
        ["UI/UX", "Homepage refinement, stronger CTA paths, proof blocks, service-page hierarchy"],
        ["AEO", "FAQs, answer sections, comparison content, process explainers"],
        ["Content", "Keyword clusters, authority articles, niche buying guides, location/category pages"],
        ["Analytics", "Lead tracking, source breakdown, CTA events, commercial dashboards"],
    ]
    add_simple_table(doc, ["Workstream", "Planned outputs"], deliverables, widths=[1.4, 4.9])

    add_title(doc, "15. Commercial Proposal", level=2)
    add_para(
        doc,
        "Kerala market signals reviewed on 6 July 2026 suggest that entry and mid-market digital services are often quoted aggressively, especially when scope is limited to SEO, basic AEO-style content structuring, and selected UI/UX refinements. Since SMO is excluded and the scope is now narrower, the commercial recommendation below has been compressed into one affordable execution package.",
        size=10.5,
        color=MUTED,
        after=8,
    )
    pricing = [
        ["Single Package", "SEO cleanup, AEO improvements, and priority UI/UX refinements across the 4 sites", "₹80,000", "Lean implementation package designed to stay within the requested cap"],
    ]
    add_simple_table(doc, ["Package", "Scope", "Fee", "Positioning"], pricing, widths=[1.1, 3.1, 0.8, 1.5])
    add_para(doc, "Recommended payment terms: 50% advance, 25% after core technical/UI updates, 25% after delivery and review.", size=10.5, color=MUTED, after=4)
    add_para(doc, "No SMO, paid ads, video production, or large-scale redevelopment is included in this package.", size=10.5, color=MUTED, after=8)
    add_para(doc, "Optional recurring support after implementation:", bold=True, after=4)
    recurring = [
        ["Monthly recurring support", "Ongoing SEO monitoring, small AEO updates, minor UI/UX refinements, reporting, and maintenance support", "₹10,000 per site / month", "12 months optional"],
        ["Total if all 4 sites are covered", "4 sites x ₹10,000 per month", "₹40,000 / month", "₹4.80L for 12 months"],
    ]
    add_simple_table(doc, ["Option", "Scope", "Rate", "Term"], recurring, widths=[1.5, 2.8, 1.0, 1.2])

    add_title(doc, "16. Package Inclusions", level=2)
    add_bullets(
        doc,
        [
            "Homepage and key-page SEO corrections on all 4 websites.",
            "Heading hierarchy cleanup on priority pages.",
            "Schema implementation on main commercial templates.",
            "FAQ and answer-structured content improvements on high-intent pages.",
            "Priority UI/UX refinements for homepage clarity, CTA paths, and trust blocks.",
            "Thin-page/index-quality review for the most important URLs.",
            "Basic reporting sheet for implemented changes and next-step recommendations.",
        ],
    )
    add_para(doc, "What is intentionally excluded to stay within budget:", bold=True, after=4)
    add_bullets(
        doc,
        [
            "Ongoing monthly SEO retainer work",
            "Social media management and posting",
            "Paid advertising management",
            "Major redesign or complete redevelopment",
            "Large-scale content production across all internal pages",
        ],
    )

    add_title(doc, "17. Final Recommendation", level=2)
    add_para(
        doc,
        "Given the ₹80,000 cap, the smartest move is a focused implementation sprint, not a broad marketing engagement. The package should concentrate on the highest-value fixes: technical SEO cleanup, AEO-ready page structuring, and practical UI/UX refinements on the most commercially important pages. This will give the client a solid base without overpromising on budget.",
        after=8,
    )

    add_title(doc, "18. Source Notes", level=2)
    add_bullets(
        doc,
        [
            "Primary domains reviewed: https://chamsglobal.com/, https://www.ksrtcbusbranding.com/, https://letsplayoutdoors.in/, https://misbahsalam.com/.",
            "Internal-page samples were selected from sitemap outputs and visible live pages on 6 July 2026.",
            "Search-footprint checks used visible branded and `site:` query results available at review time.",
            "Technical findings are based on observable frontend signals and should be complemented by server-side access if a full implementation sprint begins.",
        ],
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build()
