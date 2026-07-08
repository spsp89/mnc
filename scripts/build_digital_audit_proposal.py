from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "proposal_assets"
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "digital-audit-proposal-chams-ecosystem.docx"
RADAR_CHART = OUT_DIR / "audit-scorecard.png"
ROI_CHART = OUT_DIR / "roi-scenarios.png"
ROADMAP_CHART = OUT_DIR / "roadmap-90-days.png"


NAVY = RGBColor(11, 47, 116)
NAVY_DEEP = RGBColor(4, 28, 85)
BLUE = RGBColor(30, 92, 179)
GOLD = RGBColor(244, 178, 39)
GOLD_DEEP = RGBColor(199, 137, 8)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(97, 113, 133)
LINE = RGBColor(218, 224, 232)
SOFT = RGBColor(245, 247, 251)
SUCCESS = RGBColor(35, 138, 75)
WARN = RGBColor(190, 119, 0)
RISK = RGBColor(179, 48, 48)


@dataclass
class SiteAudit:
    name: str
    url: str
    screenshot: str
    ui_score: int
    ux_score: int
    seo_score: int
    aeo_score: int
    smo_score: int
    cro_score: int
    digital_presence: str
    uiux: list[str]
    seo: list[str]
    aeo: list[str]
    smo: list[str]
    business_potential: list[str]


SITES = [
    SiteAudit(
        name="CHAMS Global",
        url="https://chamsglobal.com/",
        screenshot="chamsglobal.png",
        ui_score=6,
        ux_score=5,
        seo_score=5,
        aeo_score=3,
        smo_score=6,
        cro_score=5,
        digital_presence="Strong master-brand positioning with blog content, social links, and multi-market messaging, but discovery and conversion are below potential.",
        uiux=[
            "Clean corporate look and good top-level service positioning.",
            "Hero area is text-heavy and low on immediate proof, resulting in slower message absorption.",
            "Primary CTA exists, but case studies, outcomes, and lead-capture points are not prominent enough."
        ],
        seo=[
            "Title and meta description are present, but Open Graph tags are missing.",
            "Canonical points to `https://www.chamsglobal.com/` while the live site resolves on non-www, which can dilute canonical clarity.",
            "Only `BreadcrumbList` schema was detected; no strong Organization, Service, FAQ, or Review schema."
        ],
        aeo=[
            "Limited structured content for answer engines to summarize.",
            "No obvious FAQ or expert-answer blocks on the homepage.",
            "Service pages need clearer problem-solution-result formatting."
        ],
        smo=[
            "Active social links are present, but some links are weak quality, including a generic Twitter home link and LinkedIn admin/feed style link.",
            "Needs stronger social proof blocks, campaign reels, and case snippets."
        ],
        business_potential=[
            "Best positioned to become the authority hub for the whole ecosystem.",
            "Can generate higher-value B2B enquiries through case-study SEO, founder content, and stronger service landing pages."
        ],
    ),
    SiteAudit(
        name="KSRTC Bus Branding",
        url="https://www.ksrtcbusbranding.com/",
        screenshot="ksrtc_retry.jpg",
        ui_score=6,
        ux_score=5,
        seo_score=4,
        aeo_score=2,
        smo_score=5,
        cro_score=6,
        digital_presence="Niche-led offer with strong category relevance and visual product clarity, but technically under-optimized and trust hygiene needs work.",
        uiux=[
            "Clear product-market fit: visitors immediately understand bus branding as the offer.",
            "Strong visual asset, but the page architecture feels dated and heavy.",
            "Quote form exists, which is good, but the trust and proof sequence can be improved."
        ],
        seo=[
            "Meta title and description exist, but no H1 was detected on the homepage.",
            "No JSON-LD schema detected.",
            "Public admin and upload directories are indexable, which is both an SEO and trust risk."
        ],
        aeo=[
            "No FAQ, pricing explainer, route coverage explainer, or campaign-format answer blocks on the homepage.",
            "No schema to help AI systems understand the business as a specialized service provider."
        ],
        smo=[
            "Social profile links exist across multiple platforms.",
            "Needs campaign walkthroughs, bus route clips, client success stories, and before/after creative posts."
        ],
        business_potential=[
            "High lead intent because the niche is specific and commercially valuable.",
            "Could rank and convert well for regional OOH and transit-media queries with better landing pages and stronger proof."
        ],
    ),
    SiteAudit(
        name="Let's Play Outdoors",
        url="https://letsplayoutdoors.in/",
        screenshot="lpo_retry.jpg",
        ui_score=7,
        ux_score=6,
        seo_score=6,
        aeo_score=4,
        smo_score=6,
        cro_score=7,
        digital_presence="Most scalable digital product concept among the four sites, with inventory-led search opportunity and transaction-style UX direction.",
        uiux=[
            "Clear marketplace proposition with searchable inventory behavior.",
            "Hero communicates intent well, but the experience can feel broad and directory-like.",
            "Needs sharper category entry points, trust markers, and conversion shortcuts."
        ],
        seo=[
            "Good title, meta description, Open Graph title/description/image, and basic Organization schema.",
            "Canonical currently points to `/home` instead of the root, which should be normalized.",
            "Large sitemap footprint creates strong programmatic SEO potential if category pages are optimized."
        ],
        aeo=[
            "FAQ exists on the domain, which is a good foundation.",
            "Needs buyer-intent explainers such as pricing, city-wise inventory, format comparison, and booking workflow content."
        ],
        smo=[
            "Brand has recognizable identity and multi-platform presence.",
            "Could perform well with inventory snapshots, campaign wins, new-city launches, and UGC-style outdoor showcases."
        ],
        business_potential=[
            "Biggest long-term upside as a searchable OOH marketplace.",
            "With category SEO and lead-routing, it can become the acquisition engine for the entire group."
        ],
    ),
    SiteAudit(
        name="Misbah Salam",
        url="https://misbahsalam.com/",
        screenshot="misbah_retry.jpg",
        ui_score=7,
        ux_score=6,
        seo_score=5,
        aeo_score=3,
        smo_score=5,
        cro_score=5,
        digital_presence="Strong founder-brand asset with personal authority potential, but content quality control and conversion framing need tightening.",
        uiux=[
            "Premium, minimal visual style with a strong personal-brand hero.",
            "Several text-spacing and copy-formatting issues reduce polish and credibility.",
            "Service and value proposition need clearer consulting packages and proof blocks."
        ],
        seo=[
            "Meta title and description exist, but Open Graph tags and schema are missing.",
            "Blog presence supports authority building, yet stronger topic clustering is needed.",
            "Needs better entity reinforcement around strategist, consultant, speaker, and mentor intents."
        ],
        aeo=[
            "Good candidate for author-led answer engine visibility through expert articles, interviews, FAQs, and topical glossaries.",
            "Current homepage lacks structured expert-answer formatting."
        ],
        smo=[
            "Good visual face for LinkedIn, YouTube Shorts, and Instagram thought leadership.",
            "Needs a more consistent publishing cadence and stronger clip-to-site funnel."
        ],
        business_potential=[
            "Can become the trust amplifier for the whole ecosystem.",
            "Founder-led authority content can improve both direct consulting leads and corporate brand lift."
        ],
    ),
]


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False):
    preferred = [
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    preferred_bold = [
        "C:/Windows/Fonts/calibrib.ttf",
        "C:/Windows/Fonts/arialbd.ttf",
    ]
    choices = preferred_bold if bold else preferred
    for path in choices:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill, outline=None, radius=24, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def make_scorecard_chart(path: Path):
    w, h = 1500, 920
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(40, bold=True)
    sub_font = font(22)
    label_font = font(20, bold=True)
    small_font = font(18)

    draw.text((60, 40), "Website Audit Scorecard", fill=hex_to_rgb("#0B2F74"), font=title_font)
    draw.text((60, 95), "Internal strategic scoring based on UI, UX, SEO, AEO, SMO, and conversion readiness", fill=hex_to_rgb("#617185"), font=sub_font)

    categories = ["UI", "UX", "SEO", "AEO", "SMO", "CRO"]
    bar_colors = ["#0B2F74", "#1D5CB3", "#F4B227", "#2E8B57"]
    max_score = 10
    chart_left = 120
    chart_top = 200
    chart_width = 1180
    group_height = 115
    bar_gap = 10
    bar_height = 16
    group_gap = 14
    label_w = 210

    for i in range(max_score + 1):
        x = chart_left + label_w + int((chart_width - label_w) * i / max_score)
        draw.line((x, chart_top - 10, x, h - 90), fill=hex_to_rgb("#E7EBF0"), width=1)
        if i < max_score:
            draw.text((x + 4, chart_top - 35), str(i), fill=hex_to_rgb("#8A96A8"), font=small_font)

    for idx, site in enumerate(SITES):
        y = chart_top + idx * (group_height + group_gap)
        draw_rounded_box(draw, (70, y - 12, w - 70, y + group_height + 10), fill=hex_to_rgb("#F8FAFD"), outline=hex_to_rgb("#E2E8F0"), radius=24)
        draw.text((95, y + 24), site.name, fill=hex_to_rgb("#111827"), font=label_font)
        scores = [site.ui_score, site.ux_score, site.seo_score, site.aeo_score, site.smo_score, site.cro_score]
        for j, (cat, score) in enumerate(zip(categories, scores)):
            yy = y + 20 + j * (bar_height + 6)
            draw.text((285, yy - 2), cat, fill=hex_to_rgb("#617185"), font=small_font)
            x0 = 335
            x1 = 335 + int((chart_width - 280) * score / max_score)
            draw.rounded_rectangle((x0, yy, w - 160, yy + bar_height), radius=8, fill=hex_to_rgb("#EEF2F7"))
            draw.rounded_rectangle((x0, yy, x1, yy + bar_height), radius=8, fill=hex_to_rgb(bar_colors[idx]))
            draw.text((w - 140, yy - 3), f"{score}/10", fill=hex_to_rgb("#111827"), font=small_font)

    img.save(path)


def make_roi_chart(path: Path):
    w, h = 1500, 920
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(40, bold=True)
    sub_font = font(22)
    label_font = font(22, bold=True)
    small_font = font(18)

    draw.text((60, 40), "90-Day ROI Scenarios", fill=hex_to_rgb("#0B2F74"), font=title_font)
    draw.text((60, 95), "Illustrative revenue upside for the combined CHAMS ecosystem after conversion, SEO, and content fixes", fill=hex_to_rgb("#617185"), font=sub_font)

    scenarios = [
        ("Conservative", 225000, "#C78908", "8 extra qualified leads/month\n1 closed deal/month"),
        ("Target", 750000, "#1D5CB3", "18 extra qualified leads/month\n2 closed deals/month"),
        ("Stretch", 1575000, "#2E8B57", "30 extra qualified leads/month\n3 closed deals/month"),
    ]
    investment = 325000

    chart_left, chart_bottom = 160, 780
    chart_top, chart_right = 180, 1380
    max_val = 1800000

    for step in range(0, max_val + 1, 300000):
        y = chart_bottom - int((chart_bottom - chart_top) * step / max_val)
        draw.line((chart_left, y, chart_right, y), fill=hex_to_rgb("#E7EBF0"), width=1)
        draw.text((60, y - 12), f"₹{step//100000}L" if step else "₹0", fill=hex_to_rgb("#8A96A8"), font=small_font)

    bar_w = 220
    gap = 110
    for idx, (name, value, color, note) in enumerate(scenarios):
        x = 260 + idx * (bar_w + gap)
        y = chart_bottom - int((chart_bottom - chart_top) * value / max_val)
        draw.rounded_rectangle((x, y, x + bar_w, chart_bottom), radius=20, fill=hex_to_rgb(color))
        draw.text((x + 25, chart_bottom + 20), name, fill=hex_to_rgb("#111827"), font=label_font)
        draw.text((x + 25, chart_bottom + 58), f"₹{value/100000:.2f}L", fill=hex_to_rgb("#111827"), font=label_font)
        draw.multiline_text((x + 25, chart_bottom + 95), note, fill=hex_to_rgb("#617185"), font=small_font, spacing=6)

    inv_y = chart_bottom - int((chart_bottom - chart_top) * investment / max_val)
    draw.line((chart_left, inv_y, chart_right, inv_y), fill=hex_to_rgb("#B33030"), width=4)
    draw.text((chart_right - 310, inv_y - 38), "Estimated project investment: ₹3.25L", fill=hex_to_rgb("#B33030"), font=label_font)

    draw_rounded_box(draw, (80, 820, 1420, 885), fill=hex_to_rgb("#F8FAFD"), outline=hex_to_rgb("#E2E8F0"), radius=18)
    draw.text((105, 840), "Assumption: revenue upside comes from more qualified inbound leads, better enquiry-to-call conversion, and stronger cross-domain trust transfer.", fill=hex_to_rgb("#425466"), font=small_font)

    img.save(path)


def make_roadmap_chart(path: Path):
    w, h = 1500, 900
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(40, bold=True)
    sub_font = font(22)
    label_font = font(22, bold=True)
    small_font = font(18)

    draw.text((60, 40), "90-Day Execution Plan", fill=hex_to_rgb("#0B2F74"), font=title_font)
    draw.text((60, 95), "Priority sequencing for design, SEO, content, tracking, and lead capture improvements", fill=hex_to_rgb("#617185"), font=sub_font)

    lanes = [
        ("UI/UX & CRO", 1, 6, "#0B2F74"),
        ("Tech SEO", 1, 4, "#F4B227"),
        ("AEO & Schema", 2, 6, "#1D5CB3"),
        ("Content & SMO", 3, 12, "#2E8B57"),
        ("Analytics & Lead Ops", 1, 8, "#C78908"),
        ("Authority Campaigns", 7, 12, "#7A3E9D"),
    ]

    grid_left, grid_top = 280, 190
    col_w = 90
    row_h = 88
    weeks = [f"W{i}" for i in range(1, 13)]

    for idx, week in enumerate(weeks):
        x = grid_left + idx * col_w
        draw_rounded_box(draw, (x, 145, x + col_w - 8, 185), fill=hex_to_rgb("#EDF3FB"), outline=hex_to_rgb("#D9E2EF"), radius=12)
        draw.text((x + 24, 156), week, fill=hex_to_rgb("#0B2F74"), font=small_font)

    for row, (label, start, end, color) in enumerate(lanes):
        y = grid_top + row * row_h
        draw.text((60, y + 24), label, fill=hex_to_rgb("#111827"), font=label_font)
        for idx in range(12):
            x = grid_left + idx * col_w
            draw.rounded_rectangle((x, y + 10, x + col_w - 8, y + 54), radius=12, fill=hex_to_rgb("#F6F8FB"), outline=hex_to_rgb("#E4EAF2"))
        x0 = grid_left + (start - 1) * col_w
        x1 = grid_left + end * col_w - 8
        draw.rounded_rectangle((x0, y + 10, x1, y + 54), radius=16, fill=hex_to_rgb(color))
        draw.text((x1 + 16, y + 20), f"Weeks {start}-{end}", fill=hex_to_rgb("#617185"), font=small_font)

    draw_rounded_box(draw, (60, 760, 1420, 855), fill=hex_to_rgb("#F8FAFD"), outline=hex_to_rgb("#E2E8F0"), radius=18)
    draw.text((90, 785), "Outcome target by Day 90: cleaner messaging, stronger local-search pages, answer-engine friendly content blocks, lead routing, dashboards, and a visible founder authority engine.", fill=hex_to_rgb("#425466"), font=small_font)
    img.save(path)


def set_cell_text(cell, text, bold=False, color=INK, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def add_title_paragraph(doc: Document, text: str, size=24, after=6, align=WD_ALIGN_PARAGRAPH.LEFT, color=NAVY_DEEP):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    return p


def add_body(doc: Document, text: str, size=11, color=INK, bold=False, italic=False, after=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    return p


def add_bullets(doc: Document, items: Iterable[str], color=INK):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        r.font.size = Pt(10.5)
        r.font.color.rgb = color
        r.font.name = "Calibri"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_numbered(doc: Document, items: Iterable[str], color=INK):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        r.font.size = Pt(10.5)
        r.font.color.rgb = color
        r.font.name = "Calibri"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p_par = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2EF")
    p_bdr.append(bottom)
    p_par.append(p_bdr)


def build_document():
    make_scorecard_chart(RADAR_CHART)
    make_roi_chart(ROI_CHART)
    make_roadmap_chart(ROADMAP_CHART)

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for section in doc.sections:
        set_page_margins(section)

    # Cover
    add_body(doc, "Prepared for CHAMS ecosystem websites", size=11, color=MUTED, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_title_paragraph(doc, "Digital Presence, SEO, AEO, SMO and 90-Day Growth Proposal", size=25, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body(doc, "Audit date: 6 July 2026 | Websites reviewed: chamsglobal.com, ksrtcbusbranding.com, letsplayoutdoors.in, misbahsalam.com", size=11, color=MUTED, after=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body(
        doc,
        "This proposal evaluates the live user experience, search readiness, social positioning, answer engine readiness, and near-term business potential of the four websites, then outlines an affordable 90-day plan to improve visibility, enquiries, and brand authority.",
        size=11,
        color=INK,
        after=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_picture(str(ASSET_DIR / "chamsglobal.png"), width=Inches(5.15))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "Prepared as a practical improvement plan for the Indian market, balancing brand quality with affordable execution.", size=10.2, color=MUTED, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Executive summary
    doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_margins(doc.sections[-1])
    add_title_paragraph(doc, "1. Executive Summary", size=18, after=8)
    add_body(
        doc,
        "The four websites already form a useful digital ecosystem: CHAMS Global is the master brand, KSRTC Bus Branding is a high-intent niche offer, Let’s Play Outdoors has strong marketplace upside, and Misbah Salam can become the authority-led trust engine. The main problem is not the lack of intent. It is fragmentation. The sites are not yet working together as a unified demand-generation system.",
        after=8,
    )
    add_body(
        doc,
        "Across the ecosystem, the strongest opportunities are better lead capture, stronger service-page clarity, structured schema for answer engines, tighter on-page copy, improved cross-linking, and a more deliberate social-content funnel. The biggest risks are technical inconsistency, weak Open Graph setup, limited answer-engine formatting, and public directory exposure on the KSRTC domain.",
        after=8,
    )
    add_body(doc, "Overall strategic read:", bold=True, after=4)
    add_bullets(
        doc,
        [
            "CHAMS Global should become the authority and case-study hub.",
            "KSRTC Bus Branding should be optimized as a high-conversion niche landing system.",
            "Let’s Play Outdoors should be scaled as the searchable inventory and lead-routing engine.",
            "Misbah Salam should be developed as the founder-expert layer that strengthens trust across all sites.",
        ],
    )

    add_title_paragraph(doc, "2. Scorecard Snapshot", size=16, after=8)
    doc.add_picture(str(RADAR_CHART), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "Scores are internal audit scores out of 10 and are meant to guide priority, not to mimic third-party tools.", size=9.5, color=MUTED, after=8)

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Website", "UI", "UX", "SEO", "AEO", "SMO", "CRO"]
    for cell, head in zip(table.rows[0].cells, headers):
        shade_cell(cell, "EAF1FB")
        set_cell_text(cell, head, bold=True, color=NAVY, size=10)
    for site in SITES:
        row = table.add_row().cells
        values = [
            site.name,
            str(site.ui_score),
            str(site.ux_score),
            str(site.seo_score),
            str(site.aeo_score),
            str(site.smo_score),
            str(site.cro_score),
        ]
        for cell, value in zip(row, values):
            set_cell_text(cell, value, size=10)
    add_body(doc, "Interpretation: Let’s Play Outdoors has the strongest scale upside, CHAMS Global is the best authority hub candidate, KSRTC Bus Branding has the clearest niche intent, and Misbah Salam can strengthen trust and thought leadership across the group.", size=10, color=MUTED, after=6)

    # Site pages
    for idx, site in enumerate(SITES, start=3):
        doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_margins(doc.sections[-1])
        add_title_paragraph(doc, f"{idx}. {site.name}", size=18, after=4)
        add_body(doc, site.url, size=10.5, color=BLUE, after=10)
        add_body(doc, site.digital_presence, size=11, color=INK, italic=True, after=10)
        doc.add_picture(str(ASSET_DIR / site.screenshot), width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_body(doc, "UI/UX Findings", bold=True, color=NAVY, after=4)
        add_bullets(doc, site.uiux)
        add_body(doc, "SEO Findings", bold=True, color=NAVY, after=4)
        add_bullets(doc, site.seo)
        add_body(doc, "AEO Findings", bold=True, color=NAVY, after=4)
        add_bullets(doc, site.aeo)
        add_body(doc, "SMO Findings", bold=True, color=NAVY, after=4)
        add_bullets(doc, site.smo)
        add_body(doc, "Business Potential", bold=True, color=NAVY, after=4)
        add_bullets(doc, site.business_potential)

    # Cross-domain findings
    doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_margins(doc.sections[-1])
    add_title_paragraph(doc, "7. Cross-Domain Strategic Findings", size=18, after=8)
    add_bullets(
        doc,
        [
            "The four websites should behave like one commercial ecosystem, but today the cross-domain funnel is weak and inconsistent.",
            "Open Graph coverage is incomplete, which reduces social sharing quality and click-through potential.",
            "Structured data is too thin for answer engines. CHAMS shows only BreadcrumbList, Let’s Play Outdoors shows Organization schema, and the other two sites showed none on the homepage.",
            "Security and trust hygiene need improvement. Standard hardening headers were not detected across the sites, and the KSRTC domain exposes public admin upload directories.",
            "Copy polish is uneven. Misbah Salam especially needs tighter editing because visible spacing and formatting issues reduce premium perception.",
            "There is blog and social activity in the ecosystem, which means the growth foundation exists. The missing layer is prioritization, repackaging, and consistent authority signaling."
        ],
    )

    add_title_paragraph(doc, "8. 90-Day Plan", size=16, after=8)
    doc.add_picture(str(ROADMAP_CHART), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "Recommended sprint structure:", bold=True, after=4)
    add_numbered(
        doc,
        [
            "Days 1-30: fix technical hygiene, improve page structure, clean copy, set analytics, normalize canonicals, add stronger CTA placements, and block exposed directories from search where required.",
            "Days 31-60: add schema, FAQ blocks, service landing pages, local SEO pages, case-study modules, and social proof. Launch founder-led authority content and niche landing enhancements.",
            "Days 61-90: expand content clusters, strengthen cross-linking, push SMO campaigns, publish video snippets and case stories, measure enquiry quality, and refine pages based on actual lead behavior."
        ],
    )

    add_title_paragraph(doc, "9. KPI Targets for the First 90 Days", size=16, after=8)
    add_bullets(
        doc,
        [
            "25% to 60% lift in qualified enquiries across the ecosystem, depending on traffic quality and sales follow-up.",
            "Improved social click quality through better Open Graph assets and more campaign-specific landing pages.",
            "Increase in local search visibility for agency, outdoor advertising, bus branding, strategist, consultant, and city-intent keywords.",
            "Better answer-engine pickup through FAQs, entity reinforcement, author content, and service-result formatting.",
            "A stronger trust narrative through case studies, founder visibility, testimonial framing, and clearer contact pathways."
        ],
    )

    add_title_paragraph(doc, "10. Commercial Proposal", size=16, after=8)
    add_body(doc, "Affordable Indian-market pricing options for the combined four-site ecosystem:", after=8)
    pricing = doc.add_table(rows=1, cols=5)
    pricing.alignment = WD_TABLE_ALIGNMENT.CENTER
    pricing.style = "Table Grid"
    price_headers = ["Plan", "Best for", "Included focus", "90-day fee", "Expected outcome"]
    for cell, head in zip(pricing.rows[0].cells, price_headers):
        shade_cell(cell, "EAF1FB")
        set_cell_text(cell, head, bold=True, color=NAVY, size=10)
    plans = [
        ["Starter", "Fix the basics fast", "UI clean-up, technical SEO fixes, CTA improvements, analytics, social asset cleanup", "₹2.25L", "Better credibility and stronger enquiry readiness"],
        ["Growth", "Recommended", "Everything in Starter plus schema, landing pages, content clustering, SMO plan, authority pages, dashboarding", "₹3.25L", "Best balance of visibility, conversion, and authority"],
        ["Authority", "Aggressive brand push", "Everything in Growth plus deeper content production, founder campaign rollout, extra city/category pages, CRO testing", "₹4.75L", "Faster growth and stronger category ownership"],
    ]
    for plan in plans:
        row = pricing.add_row().cells
        for i, value in enumerate(plan):
            set_cell_text(row[i], value, bold=(i == 0), size=9.8)

    add_body(doc, "Suggested execution model:", bold=True, after=4)
    add_bullets(
        doc,
        [
            "One-time 90-day sprint fee from the table above.",
            "Optional content and growth retainer after Day 90: ₹45,000 to ₹90,000 per month depending on deliverable volume.",
            "Paid ads, influencer spends, video shoots, and major redevelopment costs are best treated as separate line items."
        ],
    )

    add_title_paragraph(doc, "11. ROI Estimate", size=16, after=8)
    doc.add_picture(str(ROI_CHART), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(
        doc,
        "The recommended Growth plan at ₹3.25L is commercially reasonable if the ecosystem turns even a modest increase in qualified inbound traffic into 1 to 2 additional monthly closures. Because CHAMS and KSRTC appear to sell higher-value services, ROI can become attractive quickly once lead quality and trust improve.",
        after=8,
    )
    add_body(doc, "Illustrative assumptions behind the chart:", bold=True, after=4)
    add_bullets(
        doc,
        [
            "Conservative case assumes one additional closed deal per month after lead quality improves.",
            "Target case assumes two additional closed deals per month from better enquiry volume and conversion.",
            "Stretch case assumes the ecosystem starts compounding through founder authority, niche landing pages, and marketplace SEO."
        ],
    )

    add_title_paragraph(doc, "12. Immediate Recommendations", size=16, after=8)
    add_numbered(
        doc,
        [
            "Fix technical hygiene first: canonical normalization, OG tags, H1 hierarchy, schema, and security/search exposure issues.",
            "Rebuild the conversion narrative next: stronger proof, clearer offers, more visible CTAs, and better contact flows.",
            "Publish answer-engine friendly content blocks: FAQs, pricing explainers, process pages, city/category pages, and expert commentary.",
            "Use Misbah Salam as the authority bridge between brand strategy and service delivery across the ecosystem.",
            "Treat Let’s Play Outdoors as the long-term scalable acquisition engine and CHAMS Global as the central trust and case-study hub."
        ],
    )

    add_title_paragraph(doc, "13. Sources and Audit Notes", size=16, after=8)
    add_bullets(
        doc,
        [
            "Live site review conducted on 6 July 2026.",
            "Primary sources: https://chamsglobal.com/, https://www.ksrtcbusbranding.com/, https://letsplayoutdoors.in/, https://misbahsalam.com/.",
            "Technical checks included homepage metadata, canonical tags, robots/sitemap availability, schema detection, social-link presence, H1/H2 structure, and crawl/index exposure signals.",
            "Search-footprint observations included indexed pages, blog presence, and directly visible search-result snippets.",
            "ROI projections are strategic estimates and should be refined further using actual baseline traffic, close rates, and average deal value."
        ],
    )

    add_horizontal_rule(doc)
    add_body(doc, "Prepared by Codex as a practical proposal document for stakeholder review.", size=9.5, color=MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
